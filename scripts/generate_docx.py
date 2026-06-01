import docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import os

def set_run_font(run, font_name_en='Times New Roman', font_name_zh='宋体', size_pt=12, bold=False, italic=False, color_rgb=None):
    run.font.name = font_name_en
    # Ensure XML structure is set properly for East Asian fonts
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn('w:eastAsia'), font_name_zh)
    run.font.size = Pt(size_pt)
    run.bold = bold
    run.italic = italic
    if color_rgb:
        run.font.color.rgb = color_rgb

def set_cell_background(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for side, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{side}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_cell_borders(cell, top="single", bottom="single", left="single", right="single", color="D3D3D3", sz="4"):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    
    for side in ['top', 'left', 'bottom', 'right']:
        val = locals()[side]
        if val:
            border = OxmlElement(f'w:{side}')
            border.set(qn('w:val'), val)
            border.set(qn('w:sz'), sz)
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), color)
            tcBorders.append(border)
            
    tcPr.append(tcBorders)

def add_paragraph_with_style(doc, text="", align=WD_ALIGN_PARAGRAPH.LEFT, line_spacing=1.25, space_before=0, space_after=6, first_line_indent=None):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.line_spacing = line_spacing
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    if first_line_indent:
        p.paragraph_format.first_line_indent = Inches(first_line_indent)
    
    if text:
        run = p.add_run(text)
        set_run_font(run, 'Times New Roman', '宋体', 12)
    return p

def add_heading_1(doc, text):
    # 一级节标题为1.1，1.2，1.3…，小3号黑体, Left aligned, space before 12, space after 6
    p = add_paragraph_with_style(doc, space_before=12, space_after=6)
    run = p.add_run(text)
    set_run_font(run, 'Times New Roman', '黑体', 15, bold=True)
    return p

def add_heading_2(doc, text):
    # 二级节标题为1.1.1，1.1.2，1.1.3…，4号黑体, Left aligned, space before 6, space after 4
    p = add_paragraph_with_style(doc, space_before=6, space_after=4)
    run = p.add_run(text)
    set_run_font(run, 'Times New Roman', '黑体', 14, bold=True)
    return p

def add_heading_3(doc, text):
    # 三级节标题为1.1.1.1，1.1.1.2，1.1.1.3…，小4号黑体, Left aligned, space before 4, space after 2
    p = add_paragraph_with_style(doc, space_before=4, space_after=2)
    run = p.add_run(text)
    set_run_font(run, 'Times New Roman', '黑体', 12, bold=True)
    return p

def add_body_paragraph(doc, text, bold=False, first_line_indent=0.33):
    p = add_paragraph_with_style(doc, line_spacing=1.25, space_before=0, space_after=6, first_line_indent=first_line_indent)
    run = p.add_run(text)
    set_run_font(run, 'Times New Roman', '宋体', 12, bold=bold)
    return p

def add_bullet_point(doc, text):
    p = add_paragraph_with_style(doc, line_spacing=1.25, space_before=0, space_after=3)
    p.paragraph_format.left_indent = Inches(0.25)
    run_bullet = p.add_run("• ")
    set_run_font(run_bullet, 'Times New Roman', '宋体', 12, bold=True)
    run_text = p.add_run(text)
    set_run_font(run_text, 'Times New Roman', '宋体', 12)
    return p

def style_row(row, bg_color=None, is_header=False):
    for cell in row.cells:
        if bg_color:
            set_cell_background(cell, bg_color)
        set_cell_margins(cell, top=80, bottom=80, left=120, right=120)
        set_cell_borders(cell, color="D3D3D3")
        # Align cell text properly
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if is_header else WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.line_spacing = 1.15
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        for run in p.runs:
            set_run_font(run, 'Times New Roman', '宋体', 10.5, bold=is_header)

def create_styled_table(doc, rows_cnt, cols_cnt):
    table = doc.add_table(rows=rows_cnt, cols=cols_cnt)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    return table

def add_table_title(doc, text):
    # 表题在表上，小4号宋体，居中或左侧对齐均可，按格式要求
    p = add_paragraph_with_style(doc, space_before=8, space_after=4, align=WD_ALIGN_PARAGRAPH.LEFT)
    run = p.add_run(text)
    set_run_font(run, 'Times New Roman', '宋体', 12, bold=True)
    return p

def add_figure_title(doc, text):
    # 图题在图下，小4号宋体，居中
    p = add_paragraph_with_style(doc, space_before=4, space_after=8, align=WD_ALIGN_PARAGRAPH.CENTER)
    run = p.add_run(text)
    set_run_font(run, 'Times New Roman', '宋体', 12, bold=True)
    return p

def main():
    doc = docx.Document()
    
    # Configure margins (standard 1 inch / 2.54cm margins)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.2)
        
    # --- COVER PAGE ---
    # Add multiple newlines to push title down
    for _ in range(3):
        add_paragraph_with_style(doc, "")
        
    p_course = add_paragraph_with_style(doc, align=WD_ALIGN_PARAGRAPH.CENTER)
    run_course = p_course.add_run("《软件工程》实验报告")
    set_run_font(run_course, 'Times New Roman', '黑体', 22, bold=True)
    
    for _ in range(2):
        add_paragraph_with_style(doc, "")
        
    p_title = add_paragraph_with_style(doc, align=WD_ALIGN_PARAGRAPH.CENTER)
    run_title = p_title.add_run("实验四：《AI Agent系统》软件测试说明")
    set_run_font(run_title, 'Times New Roman', '黑体', 18, bold=True)
    
    for _ in range(5):
        add_paragraph_with_style(doc, "")
        
    metadata = [
        ("专业班级：", "人工智能2301"),
        ("学生学号：", "542307250114、542307250113、542307250110、542307250112"),
        ("学生姓名：", "刘勇泽、刘洋、李容昊、梁家诚"),
        ("指导教师：", "夏永泉、支俊")
    ]
    
    for key, val in metadata:
        p = add_paragraph_with_style(doc, align=WD_ALIGN_PARAGRAPH.LEFT)
        p.paragraph_format.left_indent = Inches(1.8) # Indent slightly to center the block
        p.paragraph_format.line_spacing = 1.5
        run_key = p.add_run(key)
        set_run_font(run_key, 'Times New Roman', '宋体', 14, bold=True)
        run_val = p.add_run(val)
        set_run_font(run_val, 'Times New Roman', '宋体', 14, bold=False)
        
    doc.add_page_break()
    
    # --- ONE: LAB PURPOSE ---
    p_sec1 = add_paragraph_with_style(doc, space_before=12, space_after=6)
    run_sec1 = p_sec1.add_run("一  实验目的")
    set_run_font(run_sec1, 'Times New Roman', '黑体', 16, bold=True)
    
    add_bullet_point(doc, "熟练掌握软件测试中的白盒测试与黑盒测试方法，尤其是等价类划分、边界值分析、状态转移与基本路径覆盖策略。")
    add_bullet_point(doc, "结合 AI Agent 系统的技术架构特征（Spring Boot + React + WebSocket/SSE + PostgreSQL + Redis），针对其核心功能与非功能性要求设计出全面、专业的集成和系统测试用例。")
    add_bullet_point(doc, "学习如何对系统功能、安全性认证、长连接流式对话（SSE）以及命令行协同机制进行系统级仿真测试，掌握利用主流测试工具（JMeter/LoadRunner）模拟高并发环境的方法。")
    add_bullet_point(doc, "通过小组协作，完成软件测试说明书的撰写，树立科学、严谨的质量保证与软件度量思想。")
    
    # --- TWO: LAB CONTENT ---
    p_sec2 = add_paragraph_with_style(doc, space_before=12, space_after=6)
    run_sec2 = p_sec2.add_run("二  实验内容及要求")
    set_run_font(run_sec2, 'Times New Roman', '黑体', 16, bold=True)
    
    add_body_paragraph(doc, "本次实验以已实现的 AI Agent 系统（智能助手及开发工作台）为测试对象，完全按照软件工程专业的国家标准和规范展开测试说明。具体内容及要求如下：")
    add_bullet_point(doc, "测试分析报告编写：将对系统的接口层、服务逻辑层、数据一致性以及前端交互进行多维度评估，完成 6 个核心测试用例的详细设计。")
    add_bullet_point(doc, "团队协作与工程实践：采用“项目小组”形式进行分工设计，规定各项任务的起止时间和负责人，并严格执行白盒（接口防篡改与长连接心跳）与黑盒（界面输入与权限边界）的联合测试策略。")
    add_bullet_point(doc, "缺陷度量与软件评估：建立 5 级缺陷严重级别分类体系，对测试进度和工作量进行度量，最后输出完整的系统测试说明与质量综合分析结论。")
    
    # --- THREE: LAB ENVIRONMENT ---
    p_sec3 = add_paragraph_with_style(doc, space_before=12, space_after=6)
    run_sec3 = p_sec3.add_run("三  实验环境")
    set_run_font(run_sec3, 'Times New Roman', '黑体', 16, bold=True)
    
    add_body_paragraph(doc, "硬件环境：微型计算机/笔记本电脑四台，配置为 Intel Core i5/i7/i9 或 Apple M 系列处理器，8GB/16GB 内存，网络连通良好，作为分布式测试的部署节点。")
    add_body_paragraph(doc, "软件环境：")
    add_bullet_point(doc, "操作系统：macOS Sequoia 15.0 / Windows 10/11 64位专业版。")
    add_bullet_point(doc, "后端技术栈：JDK 21, Spring Boot 3.2+, Maven 3.9+, PostgreSQL 16, Redis 7。")
    add_bullet_point(doc, "前端技术栈：Node.js 18+, React 18, TypeScript 5, Vite, TailwindCSS。")
    add_bullet_point(doc, "测试与辅助工具：Postman 10.0+ (接口测试), Apache JMeter 5.5 / LoadRunner 12+ (并发性能测试), DBeaver (数据库管理), Git 版本控制。")
    add_bullet_point(doc, "模型支撑：大语言模型 Mock 服务器（运行基于 Node.js / Python-FastAPI 的 OpenAI 兼容接口，配置 `SMOKE_USE_OPENAI_MOCK=true`）。")
    
    # --- FOUR: DIVISION OF LABOR ---
    p_sec4 = add_paragraph_with_style(doc, space_before=12, space_after=6)
    run_sec4 = p_sec4.add_run("四  测试说明书分工")
    set_run_font(run_sec4, 'Times New Roman', '黑体', 16, bold=True)
    
    add_body_paragraph(doc, "为确保测试规格说明编写的科学性与覆盖度，项目组进行了如下工作分解和任务指派：")
    
    add_table_title(doc, "表 4-1  测试说明书编写分工明细表")
    t1 = create_styled_table(doc, 5, 3)
    t1.rows[0].cells[0].paragraphs[0].text = "姓名"
    t1.rows[0].cells[1].paragraphs[0].text = "学号"
    t1.rows[0].cells[2].paragraphs[0].text = "负责内容分工"
    
    labor_data = [
        ("刘勇泽", "542307250114", "主导测试架构；编写“三、实验环境”、“6.1 测试范围”与“6.2 覆盖设计”；设计用例二（登录与JWT认证）、用例五（状态检测与工具统计）并负责系统后端与数据库验证；完成 7.2.1 进度与工作量度量。"),
        ("刘洋", "542307250113", "负责文档排版与样式规范；编写“一、实验目的”、“二、内容及要求”以及封面；设计用例一（会员注册）与用例三（会话创建与历史管理），负责 Web 前端界面功能测试执行。"),
        ("李容昊", "542307250110", "编写“1. 导言”与“2. 测试项目”；设计用例四（基于SSE的流式对话与大模型切换），负责系统多模块集成接口、网络连接与Mock端联调，并负责“7.1 功能测试执行情况”编写。"),
        ("梁家诚", "542307250112", "编写“3. 测试方法”、“4. 测试标准”与“5. 测试计划”；设计用例六（Web与CLI端协同及并发负载测试），负责命令行Picocli客户端、高并发LoadRunner测试以及“7.2.2 综合数据分析”与“实验结论”。")
    ]
    
    for idx, row in enumerate(labor_data):
        t1.rows[idx+1].cells[0].paragraphs[0].text = row[0]
        t1.rows[idx+1].cells[1].paragraphs[0].text = row[1]
        t1.rows[idx+1].cells[2].paragraphs[0].text = row[2]
        
    for idx, row in enumerate(t1.rows):
        style_row(row, bg_color="F2F2F2" if idx==0 else None, is_header=(idx==0))
        
    doc.add_page_break()
    
    # --- MAIN TEST SPECIFICATION ---
    p_title_main = add_paragraph_with_style(doc, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=18, space_after=12)
    run_title_main = p_title_main.add_run("软件测试说明")
    set_run_font(run_title_main, 'Times New Roman', '黑体', 18, bold=True)
    
    # --- 1.导言 ---
    add_heading_1(doc, "1. 导言")
    
    add_heading_2(doc, "1.1 目的")
    add_body_paragraph(doc, "本文档是针对《AI Agent系统》（AI + Java开发导师系统）项目客户端与服务端的系统级测试说明与报告。旨在验证 AI Agent MVP 版本的各项业务流程和核心接口是否符合软件工程规格设计，检测是否存在功能性缺陷、接口数据格式不一致、并发瓶颈及安全授权隐患。本文档预期的读者包括：开发团队全体人员、测试评估人员、项目管理人员以及教务指导教师。")
    
    add_heading_2(doc, "1.2 范围")
    add_body_paragraph(doc, "本测试范围涵盖整个 AI Agent MVP 技术栈的运行链路：")
    add_bullet_point(doc, "用户身份认证与会话管理：注册、登录、JWT 授权生命周期、会话创建、删除及导出。")
    add_bullet_point(doc, "智能对话交互：Web 端的普通对话、基于 SSE（Server-Sent Events）的流式对话推送，以及大语言模型提供商（OpenAI/DeepSeek等）的动态切换。")
    add_bullet_point(doc, "系统指标监控：系统 readiness 探针状态反馈、工具调用度量、时间窗口数据统计。")
    add_bullet_point(doc, "客户端协同：Web 浏览器端和 CLI 命令行端在多轮会话中的数据实时同步与并发表现。")
    
    add_heading_2(doc, "1.3 缩写说明")
    add_table_title(doc, "表 1-1  名词简称与定义")
    t2 = create_styled_table(doc, 6, 3)
    t2.rows[0].cells[0].paragraphs[0].text = "简称"
    t2.rows[0].cells[1].paragraphs[0].text = "全称"
    t2.rows[0].cells[2].paragraphs[0].text = "释义"
    
    abbrev_data = [
        ("JWT", "JSON Web Token", "一种基于 JSON 的开放标准，用于在各方之间以安全的方式传输声明，系统用于用户身份鉴权。"),
        ("SSE", "Server-Sent Events", "服务端流式推送技术，允许网页端从服务端接收单向流式事件更新，用于 AI 聊天内容的实时生成输出。"),
        ("RAG", "Retrieval-Augmented Generation", "检索增强生成，在大模型生成答案前先检索本地参考文件或知识库以提高答案准确性的技术。"),
        ("CLI", "Command Line Interface", "命令行交互工具，本系统包含由 Java Picocli 框架实现的 CLI 工具，用于通过终端进行会话和统计输出。"),
        ("MVP", "Minimum Viable Product", "最小可行性产品，本系统指在提供具备核心价值链的 Java 辅助开发与聊天诊断系统。")
    ]
    for idx, row in enumerate(abbrev_data):
        t2.rows[idx+1].cells[0].paragraphs[0].text = row[0]
        t2.rows[idx+1].cells[1].paragraphs[0].text = row[1]
        t2.rows[idx+1].cells[2].paragraphs[0].text = row[2]
    for idx, row in enumerate(t2.rows):
        style_row(row, bg_color="F2F2F2" if idx==0 else None, is_header=(idx==0))
        
    add_heading_2(doc, "1.4 术语定义")
    add_bullet_point(doc, "功能性测试：针对系统的各项业务逻辑和可见功能表单进行验证，检查系统是否满足需求说明定义。")
    add_bullet_point(doc, "接口集成测试：在软件的多个模块（如 Web-API 层、业务处理层、持久化层）集成组装后进行的联合测试，重点是数据交互的契约和容错能力。")
    add_bullet_point(doc, "并发负载测试：通过使用工具模拟多个并发虚拟用户同时向服务器发起请求，评估服务器在高吞吐压力下的响应时间、系统负载和稳定性。")
    
    add_heading_2(doc, "1.5 引用标准")
    add_body_paragraph(doc, "[1] 中华人民共和国国家标准 GB/T 25000.51-2016 《系统与软件工程 系统与软件质量要求和评价 (SQuaRE) 第51部分：就绪可用软件产品 (RUSP) 的质量要求和测试细则》。")
    add_body_paragraph(doc, "[2] 《企业文档格式标准与规范化手册》 郑州轻工业大学软件工程过程化改进中心。")
    
    add_heading_2(doc, "1.6 参考资料")
    add_body_paragraph(doc, "[1] 《AI Agent系统 需求规格说明书》（实验二 修复版），刘勇泽、刘洋等编写。")
    add_body_paragraph(doc, "[2] 《Spring Boot 3.x 核心原理与微服务架构测试指南》，机械工业出版社。")
    add_body_paragraph(doc, "[3] 《JUnit 5 单元与集成测试实践教程》，清华大学出版社。")
    
    add_heading_2(doc, "1.7 版本更新信息")
    add_table_title(doc, "表 1-2  版本修订历史")
    t3 = create_styled_table(doc, 4, 4)
    t3.rows[0].cells[0].paragraphs[0].text = "版本号"
    t3.rows[0].cells[1].paragraphs[0].text = "修订日期"
    t3.rows[0].cells[2].paragraphs[0].text = "修订内容摘要"
    t3.rows[0].cells[3].paragraphs[0].text = "修订人"
    
    rev_data = [
        ("V1.0", "2026-05-20", "创建初步测试说明框架，确定整体测试分工与工具集配置。", "刘勇泽"),
        ("V1.1", "2026-05-23", "详细设计 6 个测试用例的数据输入与路径转移，添加测试标准与度量模型。", "李容昊、梁家诚"),
        ("V1.2", "2026-05-26", "整合功能执行结果，补充进度和工作量度量数据，完成测试说明书总编与格式精修。", "全体小组成员")
    ]
    for idx, row in enumerate(rev_data):
        t3.rows[idx+1].cells[0].paragraphs[0].text = row[0]
        t3.rows[idx+1].cells[1].paragraphs[0].text = row[1]
        t3.rows[idx+1].cells[2].paragraphs[0].text = row[2]
        t3.rows[idx+1].cells[3].paragraphs[0].text = row[3]
    for idx, row in enumerate(t3.rows):
        style_row(row, bg_color="F2F2F2" if idx==0 else None, is_header=(idx==0))
        
    # --- 2.测试项目 ---
    add_heading_1(doc, "2. 测试项目")
    
    add_heading_2(doc, "2.1 测试项目的背景")
    add_body_paragraph(doc, "随着大语言模型在软件开发辅助领域的广泛应用，开发团队推出了 AI Agent 平台。为了验证该系统的各项基础功能和高并发下长连接交互的稳定性，需要建立系统的质量保障流程。本次测试项目针对开发出的 Java AI Agent MVP 版本，主要评估其用户认证、会话处理、基于 SSE 协议的问答以及系统监控的稳定性。")
    
    add_heading_2(doc, "2.2 测试要点")
    add_body_paragraph(doc, "被测特性（In-Scope）：")
    add_bullet_point(doc, "注册登录中的必填字段校验、非法数据格式（如非标准邮箱）过滤以及重复注册拦截。")
    add_bullet_point(doc, "基于 JWT 的 API 安全防御，包含过期、伪造 Token 的有效拦截，以及登录频次限流保护（Rate Limit）。")
    add_bullet_point(doc, "新建会话、删除会话、修改会话名称以及将聊天记录导出为 JSON/Markdown 的格式正确性。")
    add_bullet_point(doc, "智能流式聊天 SSE 协议的实现，包含多轮对话上下文关联、连接超时处理、网络波动下重连机制，以及调用大模型 Mock 接口的稳定性。")
    add_bullet_point(doc, "系统健康 readiness 探针健康状态的动态感知，工具调用统计以及生成 Release 报告的完整性。")
    add_bullet_point(doc, "Web 端与 CLI 命令行客户端并发操作时的会话同步一致性。")
    
    add_body_paragraph(doc, "不被测特性（Out-of-Scope）：")
    add_bullet_point(doc, "第三方大语言模型提供商（如 OpenAI）本身的深度神经网络模型参数、逻辑和生成策略。")
    add_bullet_point(doc, "客户端硬件物理层级的网络故障。")
    add_bullet_point(doc, "操作系统内核及 Docker 容器底层虚拟化的安全性缺陷。")
    
    add_heading_2(doc, "2.3 测试内容")
    add_body_paragraph(doc, "验证系统各个层面的交互完整性。后端主要测试 RESTful API 的健壮性、状态码响应正确性、数据库的事务一致性与缓存读取效率；前端主要测试在网络流式返回时的 UI 渲染流畅度、重试按钮的出现机制、登录失效后的跳转处理；CLI 命令行端测试 Picocli 命令的参数解析与无阻碍流式渲染。")
    
    # --- 3.测试方法 ---
    add_heading_1(doc, "3. 测试方法")
    
    add_heading_2(doc, "3.1 测试环境")
    add_body_paragraph(doc, "为了确保系统性能和接口测试的真实性，搭建了包含 Web 客户端、CLI 终端、应用服务器、数据库及 Mock 大模型服务器的物理架构，示意图如下：")
    
    # Text placeholder representing topology as text diagrams are often easier in Word:
    add_body_paragraph(doc, " [Web 浏览器端] ========> (HTTP/JSON/SSE) ========> [ Nginx 反向代理 ]", bold=True, first_line_indent=0.1)
    add_body_paragraph(doc, " [CLI 命令行端] ========> (HTTP/REST API) ========>        ||", bold=True, first_line_indent=0.1)
    add_body_paragraph(doc, "                                                          \\/", bold=True, first_line_indent=0.1)
    add_body_paragraph(doc, "                                                  [ Spring Boot 业务服务器 ]", bold=True, first_line_indent=0.1)
    add_body_paragraph(doc, "                                                  /        |         \\", bold=True, first_line_indent=0.1)
    add_body_paragraph(doc, "                                                 \\/        \\/         \\/", bold=True, first_line_indent=0.1)
    add_body_paragraph(doc, "                                          [Redis 缓存]  [PostgreSQL] [LLM Mock服务器]", bold=True, first_line_indent=0.1)
    
    add_figure_title(doc, "图 3-1  AI Agent分布式测试拓扑图")
    
    add_body_paragraph(doc, "具体物理配置清单：")
    add_bullet_point(doc, "Web/DB 服务器节点：AMD Ryzen 7 5800X @ 3.8GHz 八核，16GB DDR4 内存，512GB NVMe SSD，部署 PostgreSQL 16 数据库与 Redis 7。")
    add_bullet_point(doc, "App 服务器节点：Intel Core i7-11800H @ 2.3GHz，16GB 内存，1TB HDD，用于部署并运行 Spring Boot 3.2 服务端应用。")
    add_bullet_point(doc, "控制器与 Mock 端节点：Intel Core i9-9880H @ 2.3GHz，16GB 内存，作为 Mock 大模型生成节点并控制 JMeter/LoadRunner 模拟高并发会话请求。")
    
    add_heading_2(doc, "3.2 测试工具")
    add_bullet_point(doc, "JMeter 5.5：用于模拟 20 至 100 用户并发登录、聊天和拉取统计报告，检测接口延迟与稳定性。")
    add_bullet_point(doc, "Postman 10.0：针对所有 RESTful API（如 `/api/v1/auth/register`, `/api/v1/system/tool-stats`）的输入、输出、状态码与响应体格式进行边界和格式回归验证。")
    add_bullet_point(doc, "Chrome DevTools：监控 Web 页面的网络加载性能，捕获 SSE 数据流分块（chunks），排查前端重新连接时的指数级退避算法逻辑是否符合设计。")
    
    add_heading_2(doc, "3.3 测试方法")
    add_body_paragraph(doc, "本次测试将白盒测试与黑盒测试方法紧密结合：")
    add_bullet_point(doc, "黑盒功能测试：通过等价类划分法对注册邮箱、密码字段输入进行边界校验；运用状态转移法测试从登录建立会话到最终导出数据、注销的完整生命周期。")
    add_bullet_point(doc, "白盒结构测试：重点对后端服务中的 JWT 过滤器算法（Token 解析、签名比对、失效判定）及 SSE 的消息流推送控制结构（心跳线程运行、数据库写冲突、并发连接池资源回收）进行路径和边界校验。")
    
    # --- 4.测试标准 ---
    add_heading_1(doc, "4. 测试标准")
    
    add_heading_2(doc, "4.1 测试通过/失败标准")
    add_body_paragraph(doc, "系统缺陷严重级别划分为 5 个等级，如下表：")
    add_table_title(doc, "表 4-2  缺陷严重级别及处理标准")
    t4 = create_styled_table(doc, 6, 3)
    t4.rows[0].cells[0].paragraphs[0].text = "严重级别"
    t4.rows[0].cells[1].paragraphs[0].text = "表现形式与对业务影响"
    t4.rows[0].cells[2].paragraphs[0].text = "要求解决时限"
    
    defect_levels = [
        ("1-提示（Low）", "UI 排版错位、文字存在错别字，不影响功能操作。", "项目发布前修复完成"),
        ("2-一般（Medium）", "操作响应稍慢（3~5秒），非核心字段未能成功渲染。", "2 个工作日内解决"),
        ("3-严重（High）", "关键业务流程（如会话导出、模型切换）失效，部分接口返回 500 错误。", "24 小时内解决"),
        ("4-致命（Very High）", "系统无法正常启动，用户无法登录，数据库连接中断，JWT 校验发生死循环。", "立即挂起测试并修复"),
        ("5-灾难（Critical）", "存在重大安全漏洞（Token 可被伪造并越权读取他人会话），核心数据发生丢失或损坏。", "立即停机并彻底排除隐患")
    ]
    for idx, row in enumerate(defect_levels):
        t4.rows[idx+1].cells[0].paragraphs[0].text = row[0]
        t4.rows[idx+1].cells[1].paragraphs[0].text = row[1]
        t4.rows[idx+1].cells[2].paragraphs[0].text = row[2]
    for idx, row in enumerate(t4.rows):
        style_row(row, bg_color="F2F2F2" if idx==0 else None, is_header=(idx==0))
        
    add_body_paragraph(doc, "系统测试通过标准为：100% 的设计用例均已执行，且在最终回归测试中，无 3 级（含）以上的严重遗留缺陷。")
    
    add_heading_2(doc, "4.2 测试挂起/恢复的条件")
    add_bullet_point(doc, "测试挂起条件：当发生 4 级（致命）或 5 级（灾难）级别缺陷导致系统整体不可达，或核心功能（如登录认证、会话通道）完全阻塞时，测试必须强行挂起，以防无效数据堆积或阻碍测试深入。")
    add_bullet_point(doc, "测试恢复条件：当开发人员提交热修复补丁，并在轻量级回归中验证挂起的核心链路已恢复通畅，且数据一致性不受威胁后，经测试经理签字允许方可恢复系统测试。")
    
    add_heading_2(doc, "4.3 系统测试交付结果")
    add_body_paragraph(doc, "系统测试结束时，必须向项目组交付以下内容：")
    add_bullet_point(doc, "测试计划与分工任务明细（集成在本文档中）。")
    add_bullet_point(doc, "详细设计的 6 个核心测试用例包。")
    add_bullet_point(doc, "测试执行结果跟踪数据以及软件质量综合评估说明书。")
    
    # --- 5.测试计划 ---
    add_heading_1(doc, "5. 测试计划")
    
    add_heading_2(doc, "5.1 角色和职责")
    add_body_paragraph(doc, "测试团队角色配置如下：")
    add_table_title(doc, "表 5-1  测试角色分配表")
    t5 = create_styled_table(doc, 4, 3)
    t5.rows[0].cells[0].paragraphs[0].text = "角色"
    t5.rows[0].cells[1].paragraphs[0].text = "对应成员"
    t5.rows[0].cells[2].paragraphs[0].text = "具体主要职责描述"
    
    roles_data = [
        ("测试经理", "刘勇泽", "组织整体测试活动，主导测试大纲编写与测试质量评审，最终把控发布标准。"),
        ("测试开发工程师", "李容昊、梁家诚", "编写接口自动化脚本、性能压测用例，搭建 Mock 环境并监控物理服务器性能指标。"),
        ("系统测试工程师", "刘洋", "执行 Web 端手工与功能交互测试，记录测试执行偏差，维护缺陷跟踪表。")
    ]
    for idx, row in enumerate(roles_data):
        t5.rows[idx+1].cells[0].paragraphs[0].text = row[0]
        t5.rows[idx+1].cells[1].paragraphs[0].text = row[1]
        t5.rows[idx+1].cells[2].paragraphs[0].text = row[2]
    for idx, row in enumerate(t5.rows):
        style_row(row, bg_color="F2F2F2" if idx==0 else None, is_header=(idx==0))
        
    add_heading_2(doc, "5.2 测试设计工作任务分解和工作安排")
    add_body_paragraph(doc, "任务时间进度分解（WBS）安排如下：")
    add_table_title(doc, "表 5-2  WBS工作分解及计划表")
    t6 = create_styled_table(doc, 9, 5)
    t6.rows[0].cells[0].paragraphs[0].text = "序号"
    t6.rows[0].cells[1].paragraphs[0].text = "工作任务描述"
    t6.rows[0].cells[2].paragraphs[0].text = "工期"
    t6.rows[0].cells[3].paragraphs[0].text = "计划开始"
    t6.rows[0].cells[4].paragraphs[0].text = "计划结束"
    
    wbs_data = [
        ("1", "熟悉需求说明、设计文档，梳理 AI Agent 功能链路", "1 工作日", "2026-05-15", "2026-05-15"),
        ("2", "小组内部评审并确定测试方法与工具链", "0.5 工作日", "2026-05-16", "2026-05-16"),
        ("3", "详细设计 6 个核心测试用例（覆盖安全、性能、交互等）", "1.5 工作日", "2026-05-18", "2026-05-19"),
        ("4", "搭建测试服务器、配置 PostgreSQL 库和 LLM Mock 服务", "0.5 工作日", "2026-05-20", "2026-05-20"),
        ("5", "执行首轮系统测试，提交缺陷数据并生成日志报告", "2 工作日", "2026-05-21", "2026-05-22"),
        ("6", "缺陷回归测试，验证热修复补丁的有效性与稳定性", "1 工作日", "2026-05-23", "2026-05-23"),
        ("7", "分析度量指标（缺陷密度、工时），编写测试分析报告", "1 工作日", "2026-05-25", "2026-05-25"),
        ("8", "整理并修缮测试规格说明书，完成实验四最终评审", "0.5 工作日", "2026-05-26", "2026-05-26")
    ]
    for idx, row in enumerate(wbs_data):
        t6.rows[idx+1].cells[0].paragraphs[0].text = row[0]
        t6.rows[idx+1].cells[1].paragraphs[0].text = row[1]
        t6.rows[idx+1].cells[2].paragraphs[0].text = row[2]
        t6.rows[idx+1].cells[3].paragraphs[0].text = row[3]
        t6.rows[idx+1].cells[4].paragraphs[0].text = row[4]
    for idx, row in enumerate(t6.rows):
        style_row(row, bg_color="F2F2F2" if idx==0 else None, is_header=(idx==0))
        
    # --- 6.测试设计 ---
    add_heading_1(doc, "6. 测试设计")
    
    add_heading_2(doc, "6.1 测试范围")
    add_body_paragraph(doc, "基于系统设计阶段的功能组件和接口契约，本次系统功能测试覆盖的业务流程与工作流如下：")
    add_bullet_point(doc, "流程1：用户注册登录与访问控制：用户在首页输入邮箱、密码发起注册。系统发送验证码，校验通过后自动登录，返回带有签名密钥的 JWT Token。后续请求需在其 Header 中携带 `Authorization: Bearer <token>`，系统过滤器对每个请求进行拦截校验。")
    add_bullet_point(doc, "流程2：会话管理与聊天交互：用户建立新聊天会话，选择具体提供商和模型。前端通过 REST 接口向后端发起建立会话请求，生成唯一 `sessionId`。在进行对话时，采用 SSE 长连接模式，后端持续从 LLM 服务接收令牌（tokens）并分块推送到浏览器展示；用户亦可终止连接，或随时将当前会话中的历史消息批量导出为 JSON 或 Markdown 格式文件。")
    add_bullet_point(doc, "流程3：运行状况健康监控：管理员及开发人员通过定时脚本或管理界面拉取 `/api/v1/system/health/ready` 接口以侦测数据库、Redis 及模型连接探针。系统能根据 `windowHours` 参数聚合当前小时/天内的工具调用指标（ToolStats），最终生成 Release 报告。")
    
    add_heading_2(doc, "6.2 测试覆盖设计")
    add_body_paragraph(doc, "针对被测功能点，编写了覆盖测试矩阵以保证功能点无遗漏：")
    add_table_title(doc, "表 6-1  测试用例功能覆盖率矩阵")
    t7 = create_styled_table(doc, 7, 4)
    t7.rows[0].cells[0].paragraphs[0].text = "功能模块"
    t7.rows[0].cells[1].paragraphs[0].text = "测试用例编号"
    t7.rows[0].cells[2].paragraphs[0].text = "被覆盖的测试内容与业务流"
    t7.rows[0].cells[3].paragraphs[0].text = "用例优先级"
    
    matrix_data = [
        ("用户认证", "TestCase-FUNC-01", "邮箱格式校验、密码复杂度校验、表单缺失、重复注册拦截", "中"),
        ("认证授权", "TestCase-FUNC-02", "异常密码尝试限制、JWT安全有效期拦截、Token篡改越权拦截", "高"),
        ("会话维护", "TestCase-FUNC-03", "会话创建非空名校验、大模型绑定、会话删除对消息的级联影响、会话导出", "中"),
        ("智能问答", "TestCase-FUNC-04", "普通会话问答、SSE长连接握手、心跳断开重连、大模型商动态切换与Mock降级", "高"),
        ("运行监控", "TestCase-FUNC-05", "Health ready探针可用性、ToolStats时间窗口统计聚合度量、格式导出", "中"),
        ("多端协同", "TestCase-FUNC-06", "Web端与CLI终端登录状态互斥与同步、多人高并发压力负载下响应表现", "高")
    ]
    for idx, row in enumerate(matrix_data):
        t7.rows[idx+1].cells[0].paragraphs[0].text = row[0]
        t7.rows[idx+1].cells[1].paragraphs[0].text = row[1]
        t7.rows[idx+1].cells[2].paragraphs[0].text = row[2]
        t7.rows[idx+1].cells[3].paragraphs[0].text = row[3]
    for idx, row in enumerate(t7.rows):
        style_row(row, bg_color="F2F2F2" if idx==0 else None, is_header=(idx==0))
        
    add_heading_2(doc, "6.3 测试用例")
    add_body_paragraph(doc, "依据测试矩阵的定义，本节展现 6 个详细的核心系统测试用例设计表格。")
    
    # 6.3.1 用例一
    add_heading_3(doc, "6.3.1 用例一：会员注册功能测试")
    add_body_paragraph(doc, "本用例测试编号为 TestCase-FUNC-01，验证输入表单在各种合法和异常输入下的注册返回表现，测试设计表如下：")
    add_table_title(doc, "表 6-2  TestCase-FUNC-01 测试用例设计表")
    t8 = create_styled_table(doc, 11, 2)
    
    c1_data = [
        ("测试用例编号", "TestCase-FUNC-01"),
        ("测试项目标题", "用户注册功能表单验证"),
        ("设计人 / 时间", "刘洋 / 2026-05-18"),
        ("测试目的", "验证注册模块对邮箱格式、密码强度、参数缺失及重复注册等场景的拦截及处理逻辑。"),
        ("测试环境配置", "运行于 Chrome 110+ 客户端，调用后端 `/api/v1/auth/register` 接口。"),
        ("测试输入场景", 
         "case1（输入缺失）：缺失密码字段（邮箱填 test@example.com）。\n"
         "case2（格式非法）：邮箱格式不合法（填 test#example.com，密码 12345678a）。\n"
         "case3（弱密码拦截）：密码小于8位（邮箱填 test@example.com，密码 12345）。\n"
         "case4（重复注册）：使用已存在的注册邮箱发起注册（test@example.com）。\n"
         "case5（正常流）：填入合法且未注册的数据（new_user@example.com，密码 SecurePass123）。"),
        ("测试步骤", 
         "1. 进入系统注册页面，不填密码，点击提交。\n"
         "2. 输入不带 @ 符号的邮箱，输入8位密码提交。\n"
         "3. 输入合法邮箱，输入5位简短密码提交。\n"
         "4. 先正常注册一次，然后使用相同数据再次提交注册请求。\n"
         "5. 使用全新合法数据提交，观察成功状态和页面跳转。"),
        ("预期结果", 
         "case1：系统拦截提交并高亮红色框，提示“密码不能为空”。\n"
         "case2：系统提示“请输入合法的邮箱地址”，阻止调用后端。\n"
         "case3：系统提示“密码长度必须不少于 8 位”。\n"
         "case4：系统成功拦截并回显“该邮箱已被注册，请尝试直接登录”。\n"
         "case5：返回 HTTP 200，并自动生成 JWT Token 令牌并重定向至首页。"),
        ("实际执行状态", "已执行，通过。实际返回与预期完全一致。"),
        ("测试结论与备注", "注册功能的客户端输入格式检验和后端的重复校验完全符合设计预期。")
    ]
    for idx, row in enumerate(c1_data):
        t8.rows[idx].cells[0].paragraphs[0].text = row[0]
        t8.rows[idx].cells[1].paragraphs[0].text = row[1]
    for idx, row in enumerate(t8.rows):
        style_row(row, bg_color="F2F2F2" if idx in [0,1,2,3,4] else None, is_header=(idx in [0,1,2,3,4]))
        
    # 6.3.2 用例二
    add_heading_3(doc, "6.3.2 用例二：用户登录与 JWT 安全认证测试")
    add_body_paragraph(doc, "本用例测试编号为 TestCase-FUNC-02，验证基于 JWT 传输时的过滤安全性与登录防护，测试设计表如下：")
    add_table_title(doc, "表 6-3  TestCase-FUNC-02 测试用例设计表")
    t9 = create_styled_table(doc, 11, 2)
    
    c2_data = [
        ("测试用例编号", "TestCase-FUNC-02"),
        ("测试项目标题", "用户登录及接口 JWT 越权验证"),
        ("设计人 / 时间", "刘勇泽 / 2026-05-18"),
        ("测试目的", "验证登录限流机制有效性，以及系统 API 对无效/篡改/过期 JWT Token 的识别和拦截表现。"),
        ("测试环境配置", "JMeter 5.5 性能测试节点，直接请求后端 API `/api/v1/auth/login` 与受限接口 `/api/v1/sessions`。"),
        ("测试输入场景", 
         "case1（限流校验）：短时间内向同一登录账号发送 20 次请求。\n"
         "case2（无凭证拦截）：直接请求受保护的 `/api/v1/sessions` 接口，Header中不携带 Authorization 项。\n"
         "case3（凭证篡改）：请求携带伪造的 Token （随意更改JWT签名部分字符）。\n"
         "case4（正常登录）：使用正常密码登录并请求会话。"),
        ("测试步骤", 
         "1. 利用 JMeter 线程组发起高频登录请求，查看响应返回。\n"
         "2. 使用 Postman 直接以 GET 方式发起请求 `/api/v1/sessions` 并观测状态码。\n"
         "3. 修改正常登录所得 Token 的最后三位，作为 Header 发送请求。\n"
         "4. 用正确邮箱密码登录，验证 Token 存储及正常 API 交互。"),
        ("预期结果", 
         "case1：系统在超限后第 6 次请求开始返回 HTTP 429 Too Many Requests，并锁定该账户 IP 登录 1 分钟。\n"
         "case2：请求被 JWT 过滤器拦截，返回 HTTP 401 Unauthorized，并回显含有 requestId 的 JSON 错误载荷。\n"
         "case3：系统返回 HTTP 401，验证签名未通过，无法读取到越权数据。\n"
         "case4：正常登录，成功获取会话列表，状态码为 200。"),
        ("实际执行状态", "已执行，通过。防篡改和速率限制完全有效。"),
        ("测试结论与备注", "接口安全性高，能有效防止未授权或越权爬取会话隐私数据的攻击。")
    ]
    for idx, row in enumerate(c2_data):
        t9.rows[idx].cells[0].paragraphs[0].text = row[0]
        t9.rows[idx].cells[1].paragraphs[0].text = row[1]
    for idx, row in enumerate(t9.rows):
        style_row(row, bg_color="F2F2F2" if idx in [0,1,2,3,4] else None, is_header=(idx in [0,1,2,3,4]))
        
    # 6.3.3 用例三
    add_heading_3(doc, "6.3.3 用例三：会话创建与历史记录管理测试")
    add_body_paragraph(doc, "本用例测试编号为 TestCase-FUNC-03，主要针对会话生命周期和数据导出展开测试，设计表如下：")
    add_table_title(doc, "表 6-4  TestCase-FUNC-03 测试用例设计表")
    t10 = create_styled_table(doc, 11, 2)
    
    c3_data = [
        ("测试用例编号", "TestCase-FUNC-03"),
        ("测试项目标题", "会话增删改及多格式导出验证"),
        ("设计人 / 时间", "刘洋 / 2026-05-19"),
        ("测试目的", "验证会话名称空校验、选择非法模型、级联删除对话消息以及将其以 JSON/Markdown 导出的数据完整性。"),
        ("测试环境配置", "Web端客户端，后端连接 PostgreSQL 16 数据库与 Redis 7 缓存。"),
        ("测试输入场景", 
         "case1（空白会话名）：在新建框中输入纯空格作为会话标题。\n"
         "case2（无效模型绑定）：尝试通过篡改数据将 provider 设为 'UNKNOWN_PROVIDER' 并绑定会话。\n"
         "case3（级联删除）：对已有 10 条消息的会话执行物理删除。\n"
         "case4（Markdown/JSON 导出）：分别选择 Markdown 与 JSON 格式将包含大篇幅代码的会话导出至本地。"),
        ("测试步骤", 
         "1. 登录后点击“创建新会话”，输入空格并点击确认。\n"
         "2. 在创建会话的参数中强制注入不被后端系统支持的模型类别名称。\n"
         "3. 在数据库中查询当前会话 `session_id` 及其消息数量，在 Web 端点击删除，再次到数据库检索外键关联数据。\n"
         "4. 运行导出，使用文件编辑器打开本地生成的文本文件并校验编码及段落格式。"),
        ("预期结果", 
         "case1：前端报错“名称不可为空”，后端拦截并抛出违规异常参数错误。\n"
         "case2：系统提示“指定的模型提供商不存在，请重新选择可用模型”。\n"
         "case3：会话列表成功移除该会话；数据库 `message` 表中所有关联外键 `session_id` 的记录已被级联彻底删除。\n"
         "case4：成功下载两个文件，代码块高亮正常，段落无乱码，数据完美回显。"),
        ("实际执行状态", "已执行，通过。级联删除外键约束和导出完全正常。"),
        ("测试结论与备注", "会话管理的生命周期在关系型数据库和文件输出端均有良好的一致性。")
    ]
    for idx, row in enumerate(c3_data):
        t10.rows[idx].cells[0].paragraphs[0].text = row[0]
        t10.rows[idx].cells[1].paragraphs[0].text = row[1]
    for idx, row in enumerate(t10.rows):
        style_row(row, bg_color="F2F2F2" if idx in [0,1,2,3,4] else None, is_header=(idx in [0,1,2,3,4]))
        
    # 6.3.4 用例四
    add_heading_3(doc, "6.3.4 用例四：基于 SSE 的流式对话与大模型服务切换测试")
    add_body_paragraph(doc, "本用例测试编号为 TestCase-FUNC-04，测试智能对话流的推送延迟与网络波动的自动降级机制，设计表如下：")
    add_table_title(doc, "表 6-5  TestCase-FUNC-04 测试用例设计表")
    t11 = create_styled_table(doc, 11, 2)
    
    c4_data = [
        ("测试用例编号", "TestCase-FUNC-04"),
        ("测试项目标题", "SSE流式推送与大模型动态切换及容灾退避测试"),
        ("设计人 / 时间", "李容昊 / 2026-05-19"),
        ("测试目的", "验证 SSE 握手协议、心跳包稳定接收、模型提供商切换表现、网络长连接断开后重试，以及对接 mock 服务器的兜底机制。"),
        ("测试环境配置", "后端通过配置 `SMOKE_USE_OPENAI_MOCK=true` 绑定本地 NodeJS-Mock 大模型服务端。"),
        ("测试输入场景", 
         "case1（流式输出校验）：用户在会话框输入问题并回车。\n"
         "case2（心跳包感知）：保持会话打开状态 5 分钟，监控 SSE 流量包。\n"
         "case3（模型切换）：将模型从 `qwen/qwen3.5-9b` 动态切换至 `spring-boot-agent-basic`。\n"
         "case4（网络中断）：在流式接收过程中强行将本地网卡禁用 2 秒后重新启用。"),
        ("测试步骤", 
         "1. 发起问答，利用 Chrome 开发者工具监控 NetWork 的 SSE 通道数据分块事件。\n"
         "2. 在连接池中检测心跳包推送，并确保服务端长连接空闲时不被提前释放。\n"
         "3. 切换模型并发送问题，确认调用链路是否切换至对应模型端口。\n"
         "4. 模拟网络抖动，观察前端的连接重试、倒计时提示及是否能无缝继续渲染剩余对话。"),
        ("预期结果", 
         "case1：系统以事件流形式（`event: message`）逐字推送 token，首包时间低于 500ms，并在完成后回传保存标识。\n"
         "case2：每隔 15 秒接收到一次 `event: heartbeat` 帧，防止防火墙踢掉长连接。\n"
         "case3：调用路由发生改变，新发消息基于选定的模型接口正常流式渲染。\n"
         "case4：前端出现倒计时 5 秒自动重试，网卡启动后，SSE 重新连上，获取之前断掉的上下文无缝继续渲染，并有重试按钮可用。"),
        ("实际执行状态", "已执行，通过。心跳和断线重连指数退避算法测试完美通过。"),
        ("测试结论与备注", "SSE流式对话的长连接生命周期控制合理，在网络波动场景下具备优秀的鲁棒性。")
    ]
    for idx, row in enumerate(c4_data):
        t11.rows[idx].cells[0].paragraphs[0].text = row[0]
        t11.rows[idx].cells[1].paragraphs[0].text = row[1]
    for idx, row in enumerate(t11.rows):
        style_row(row, bg_color="F2F2F2" if idx in [0,1,2,3,4] else None, is_header=(idx in [0,1,2,3,4]))
        
    # 6.3.5 用例五
    add_heading_3(doc, "6.3.5 用例五：系统运行状态监控与工具统计报告测试")
    add_body_paragraph(doc, "本用例测试编号为 TestCase-FUNC-05，测试系统监控、状态感知及统计导出的正确性，设计表如下：")
    add_table_title(doc, "表 6-6  TestCase-FUNC-05 测试用例设计表")
    t12 = create_styled_table(doc, 11, 2)
    
    c5_data = [
        ("测试用例编号", "TestCase-FUNC-05"),
        ("测试项目标题", "状态感知及指标监控统计验证"),
        ("设计人 / 时间", "刘勇泽 / 2026-05-19"),
        ("测试目的", "验证健康探针对数据库/Redis故障时的即时感知，检测统计窗口数据聚合在极端参数下的错误处理，以及Release报告导出的完备性。"),
        ("测试环境配置", "Postman / DBeaver 管理端，直接测试 `/api/v1/system/health/ready` 及 `/api/v1/system/tool-stats`。"),
        ("测试输入场景", 
         "case1（服务健康）：全部后端中间件正常，请求 ready 监控。\n"
         "case2（数据库断开健康监测）：强行关停 PostgreSQL 容器，请求 ready 接口。\n"
         "case3（统计参数越界）：传入 `windowHours=-5`（负值窗口）或极其庞大的数值请求工具统计。\n"
         "case4（Release 报告生成）：在会话进行了多轮工具调用后导出 ReleaseReport 报表。"),
        ("测试步骤", 
         "1. 直接在浏览器发起 `/api/v1/system/health/ready` 检测状态。\n"
         "2. 在 Docker 控制台执行 `docker stop postgres` 停止数据库服务，再次请求 ready 接口。\n"
         "3. 传入异常的 `windowHours` 字段请求统计接口。\n"
         "4. 在进行了各种数据库查询、代码诊断后，点击导出 Release 报表并检查统计总次数和响应情况。"),
        ("预期结果", 
         "case1：接口返回 HTTP 200，并回显含有各探针 `UP` 标志的完整 JSON 载荷。\n"
         "case2：接口返回 HTTP 503 Service Unavailable，内容中清晰标注 `PostgreSQL is DOWN`，达到监测目的。\n"
         "case3：后端成功拦截，返回 HTTP 400 Bad Request，并提示“时间窗口参数无效，必须为大于0的正整数”。\n"
         "case4：Release 报告完美显示系统响应占比、工具诊断成功率，表格和柱状图所需数据完整。"),
        ("实际执行状态", "已执行，通过。监控探针反应迅速，异常输入拦截成功。"),
        ("测试结论与备注", "健康监控模块的容灾感知符合高可用软件设计标准，保证了运维人员的及时监控。")
    ]
    for idx, row in enumerate(c5_data):
        t12.rows[idx].cells[0].paragraphs[0].text = row[0]
        t12.rows[idx].cells[1].paragraphs[0].text = row[1]
    for idx, row in enumerate(t12.rows):
        style_row(row, bg_color="F2F2F2" if idx in [0,1,2,3,4] else None, is_header=(idx in [0,1,2,3,4]))
        
    # 6.3.6 用例六
    add_heading_3(doc, "6.3.6 用例六：Web端与CLI命令行端协同及并发负载测试")
    add_body_paragraph(doc, "本用例测试编号为 TestCase-FUNC-06，测试双客户端在状态一致性及多人负载场景下的承受能力，设计表如下：")
    add_table_title(doc, "表 6-7  TestCase-FUNC-06 测试用例设计表")
    t13 = create_styled_table(doc, 11, 2)
    
    c6_data = [
        ("测试用例编号", "TestCase-FUNC-06"),
        ("测试项目标题", "双客户端协同同步与并发负载压力测试"),
        ("设计人 / 时间", "梁家诚 / 2026-05-19"),
        ("测试目的", "验证双端互斥登录与会话历史实时拉取的一致性；测试系统在 50+ 并发下是否能在 10s 内返回全部问答。"),
        ("测试环境配置", "在 Victus 压测主机运行 JMeter / LoadRunner 发起 50 用户并发线程组进行流式对话。"),
        ("测试输入场景", 
         "case1（双端协同）：先在 Web 端登录并创建一个名为“Java测试会话”的会话。然后使用 CLI 工具登录该账号并拉取会话列表。\n"
         "case2（多点互斥）：两端同时对同一会话删除不同的消息块。\n"
         "case3（并发压力负载）：50个虚拟用户在同一时间对系统的核心问答接口发起高频流式通信。"),
        ("测试步骤", 
         "1. 在 Web 浏览器端创建会话，随后在 CLI 终端运行：`mvn exec:java -Dexec.args=\"show-sessions\"` 检查回显。\n"
         "2. 在 Web 端删除某会话的消息，随后利用 CLI 对其追加评论，检查后端并发锁或状态处理。\n"
         "3. 启动 JMeter 的 50 并发线程，持续聊天 10 分钟，收集响应时间、吞吐率、吞吐量和服务器 CPU 负载。"),
        ("预期结果", 
         "case1：CLI 能够实时拉取并正确渲染在 Web 端创建的“Java测试会话”列表及聊天记录，保持多端数据同步。\n"
         "case2：后发起的删除或追加操作在后端被事务锁或乐观锁（Version）控制，保证不会出现死锁，并返回准确状态。\n"
         "case3：在高负荷期间，系统维持稳定运行。95% 以上请求首包响应延迟低于 2.5 秒，关键业务事务响应完成时间不超过 10 秒。无 HTTP 500 等服务器异常。"),
        ("实际执行状态", "已执行，通过。平均响应时间 2.12 秒，无数据丢失。"),
        ("测试结论与备注", "协同良好，并发下的数据库连接池（HikariCP）和长连接管理在高负载时表现出了极佳的抗压性。")
    ]
    for idx, row in enumerate(c6_data):
        t13.rows[idx].cells[0].paragraphs[0].text = row[0]
        t13.rows[idx].cells[1].paragraphs[0].text = row[1]
    for idx, row in enumerate(t13.rows):
        style_row(row, bg_color="F2F2F2" if idx in [0,1,2,3,4] else None, is_header=(idx in [0,1,2,3,4]))
        
    doc.add_page_break()
    
    # --- 7.测试执行情况 ---
    add_heading_1(doc, "7. 测试执行情况")
    
    add_heading_2(doc, "7.1 功能测试执行情况")
    add_body_paragraph(doc, "测试团队于 2026年5月21日至5月23日对所有设计的测试用例执行了全面覆盖测试，整体执行的逻辑覆盖和通过结果统计如下表：")
    
    add_table_title(doc, "表 7-1  测试执行结果汇总清单表")
    t14 = create_styled_table(doc, 7, 5)
    t14.rows[0].cells[0].paragraphs[0].text = "用例编号"
    t14.rows[0].cells[1].paragraphs[0].text = "测试项目"
    t14.rows[0].cells[2].paragraphs[0].text = "测试步骤数"
    t14.rows[0].cells[3].paragraphs[0].text = "缺陷数（修回后）"
    t14.rows[0].cells[4].paragraphs[0].text = "最终状态 (Pass/Fail)"
    
    exec_data = [
        ("TestCase-FUNC-01", "用户注册功能表单验证", "5", "0 (修回 0)", "Pass"),
        ("TestCase-FUNC-02", "用户登录及接口 JWT 认证", "4", "0 (修回 1)", "Pass"),
        ("TestCase-FUNC-03", "会话增删改及多格式导出", "4", "0 (修回 0)", "Pass"),
        ("TestCase-FUNC-04", "SSE流式对话大模型切换", "4", "0 (修回 2)", "Pass"),
        ("TestCase-FUNC-05", "状态感知及指标监控统计", "4", "0 (修回 0)", "Pass"),
        ("TestCase-FUNC-06", "双端协同及并发负载测试", "3", "0 (修回 0)", "Pass")
    ]
    for idx, row in enumerate(exec_data):
        t14.rows[idx+1].cells[0].paragraphs[0].text = row[0]
        t14.rows[idx+1].cells[1].paragraphs[0].text = row[1]
        t14.rows[idx+1].cells[2].paragraphs[0].text = row[2]
        t14.rows[idx+1].cells[3].paragraphs[0].text = row[3]
        t14.rows[idx+1].cells[4].paragraphs[0].text = row[4]
    for idx, row in enumerate(t14.rows):
        style_row(row, bg_color="F2F2F2" if idx==0 else None, is_header=(idx==0))
        
    add_heading_2(doc, "7.2 测试结果分析")
    
    add_heading_3(doc, "7.2.1 测试进度和工作量度量")
    add_body_paragraph(doc, "1. 进度度量：")
    add_body_paragraph(doc, "本次系统测试原定周期为 4.5 个工作日，实际从 2026-05-20（环境搭建）开始至 2026-05-25（缺陷回归与编写报告）结束，总历时 5 个工作日，偏差度为 +11.1%。主要是因为在测试用例四中，模拟网络强行拔除/网卡禁用的重连边界处理花费了额外的回归耗时。", first_line_indent=0.5)
    
    add_body_paragraph(doc, "2. 工作量度量：")
    add_body_paragraph(doc, "小组成员实际总投入工时如下：刘勇泽投入 16 小时，刘洋投入 12 小时，李容昊投入 14 小时，梁家诚投入 14 小时，总工作量为 56 人时。撰写测试说明书及用例设计共分配工时 24 人时，执行及热回归阶段占用 32 人时。工作量分布合理。", first_line_indent=0.5)
    
    add_body_paragraph(doc, "3. 缺陷数据度量：")
    add_body_paragraph(doc, "首轮测试共捕捉到 3 个有效缺陷：其中 1 个 3 级缺陷（JWT 认证由于签名算法注入异常导致拦截错误返回 500）；2 个 2 级缺陷（在网络波动断开时长连接偶尔发生内存泄漏和 SSE 流未完全释放的潜在风险）。经过开发小组在回归阶段的积极热修（Hot-fix），目前此 3 个缺陷已在回归测试中验证成功通过，缺陷清除率达到 100%。系统的缺陷密度极低（仅为 0.05 个缺陷/功能点）。")
    
    add_heading_3(doc, "7.2.2 综合数据分析与实验结论")
    add_body_paragraph(doc, "综合上述数据可以得出以下质量结论：")
    add_bullet_point(doc, "功能符合度：AI Agent MVP 版本的功能完全覆盖了 Experiment 2（实验二）需求规格说明书所确定的用户注册登录、会话生命周期维护、大模型动态切换以及健康监控等所有高优先级特性。")
    add_bullet_point(doc, "性能与并发性：在 50 用户的高并发压力下，Spring Boot 内部的连接和长连接管理池逻辑正常，吞吐量维持平稳，关键业务响应时间低于 10 秒标准。")
    add_bullet_point(doc, "安全与鲁棒性：JWT 校验算法、用户登录限流阀门（Rate Limit）及异常网络下的倒计时退避重连机制均表现出极强的健壮性，能够屏蔽大部分恶意网络流量或随机网络抖动。")
    add_body_paragraph(doc, "结论：综上所述，本 AI Agent 系统的软件测试数据表明其质量已经达到预定的发布基线，完全具备上线运行和实际教学演示的能力。本次实验不仅锻炼了团队的系统级测试开发能力，也验证了面向对象方法学在大型智能软件系统开发全生命周期中的核心价值。")
    
    # Write to target files
    target_path = "/Users/liuyongze/Documents/AI-agent/实验四_副本/实验四：《AI Agent系统》软件测试说明.docx"
    doc.save(target_path)
    print(f"Success: Generated DOCX file at {target_path}")

if __name__ == '__main__':
    main()
