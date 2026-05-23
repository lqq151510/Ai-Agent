from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
UML_PNG_DIR = ROOT / "uml" / "png"


def set_default_font(doc: Document, font_name: str = "宋体", size_pt: int = 12) -> None:
    style = doc.styles["Normal"]
    style.font.name = font_name
    style._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    style.font.size = Pt(size_pt)


def add_center_title(doc: Document, text: str, size_pt: int = 16, bold: bool = True) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size_pt)
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")


def add_kv_line(doc: Document, k: str, v: str) -> None:
    p = doc.add_paragraph()
    p.add_run(f"{k}：").bold = True
    p.add_run(v)


def add_section_title(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True


def add_numbered_lines(doc: Document, lines: list[str]) -> None:
    for idx, line in enumerate(lines, 1):
        doc.add_paragraph(f"{idx}. {line}")


def add_term_table(doc: Document, rows: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=len(rows) + 1, cols=2)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "术语"
    hdr[1].text = "含义"
    for i, (term, meaning) in enumerate(rows, start=1):
        r = table.rows[i].cells
        r[0].text = term
        r[1].text = meaning


def add_division_table(doc: Document, rows: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=len(rows) + 1, cols=2)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "成员"
    hdr[1].text = "负责内容"
    for i, (name, work) in enumerate(rows, start=1):
        r = table.rows[i].cells
        r[0].text = name
        r[1].text = work


def add_usecase(doc: Document, title: str, actor: str, pre: str, post: str, main: list[str], alt: list[str] | None = None) -> None:
    p = doc.add_paragraph()
    r = p.add_run(title)
    r.bold = True
    doc.add_paragraph(f"参与者：{actor}")
    doc.add_paragraph(f"前置条件：{pre}")
    doc.add_paragraph(f"后置条件：{post}")
    doc.add_paragraph("主事件流：")
    for s in main:
        doc.add_paragraph(s)
    if alt:
        doc.add_paragraph("备选事件流：")
        for s in alt:
            doc.add_paragraph(s)


def add_uml_figure(doc: Document, filename: str, caption: str) -> None:
    img_path = UML_PNG_DIR / filename
    if not img_path.exists():
        raise FileNotFoundError(str(img_path))
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(img_path), width=Inches(5.8))
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER


def build_doc(out_path: Path) -> None:
    doc = Document()
    set_default_font(doc)

    # 封面
    add_center_title(doc, "《软件工程》实验报告", size_pt=18)
    doc.add_paragraph("")
    add_center_title(doc, "实验二：AI Agent系统需求规格说明", size_pt=16)
    doc.add_paragraph("")
    add_kv_line(doc, "专业班级", "人工智能2301")
    add_kv_line(doc, "学生学号", "542307250114、542307250113、542307250110、542307250112")
    add_kv_line(doc, "学生姓名", "刘勇泽、刘洋、李容昊、梁家诚")
    add_kv_line(doc, "指导教师", "夏永泉、支俊")

    doc.add_paragraph("")

    # 一 实验目的
    add_section_title(doc, "一  实验目的")
    add_numbered_lines(
        doc,
        [
            "结合实际开发完成的AI Agent系统，对系统进行完整的软件需求分析。",
            "掌握面向对象需求分析方法，并使用UML对系统功能进行建模。",
            "学习软件工程中需求规格说明书的编写规范。",
            "为后续系统设计、编码实现、系统测试与项目验收提供依据。",
            "提升从“功能实现”到“工程化文档表达”的能力。",
        ],
    )

    # 二 实验内容及要求
    add_section_title(doc, "二  实验内容及要求")
    add_numbered_lines(
        doc,
        [
            "以AI Agent系统作为实验对象，完成需求规格说明书编写。",
            "对系统功能、业务流程、运行环境及非功能需求进行分析。",
            "使用UML建立系统用例模型与功能结构模型。",
            "对系统主要业务流程进行结构化描述。",
            "输出符合软件工程课程规范的需求规格说明书。",
        ],
    )

    # 三 实验环境
    add_section_title(doc, "三  实验环境")
    doc.add_paragraph("硬件环境")
    doc.add_paragraph("微型计算机一台。")
    doc.add_paragraph("软件环境")
    doc.add_paragraph(
        "Windows 10操作系统、Word 2016、Rational Rose 2003、JDK 21、Maven 3.9+、Node.js 18+、Docker Desktop、PostgreSQL 16、Redis 7、IntelliJ IDEA、VS Code、Git。"
    )

    # 四 分工
    add_section_title(doc, "四  需求说明书分工")
    add_division_table(
        doc,
        [
            ("刘勇泽", "需求分析、后端设计、系统总体架构"),
            ("刘洋", "前端界面设计、文档排版"),
            ("李容昊", "系统测试、接口联调"),
            ("梁家诚", "UML建模、RAG（基于知识检索的增强生成技术）与Worker模块"),
        ],
    )

    doc.add_page_break()

    # 需求说明
    add_section_title(doc, "需求说明")

    add_section_title(doc, "1  概述")

    add_section_title(doc, "1.1  编写目的")
    doc.add_paragraph(
        "本文档用于描述AI Agent系统的功能需求、性能需求、运行环境及系统约束，作为系统开发、测试、维护与验收的重要依据。"
    )
    doc.add_paragraph("通过编写本需求规格说明书，可以：")
    doc.add_paragraph("（1）明确系统建设目标与业务需求；")
    doc.add_paragraph("（2）统一项目成员对系统功能与流程的理解；")
    doc.add_paragraph("（3）为后续系统设计与编码实现提供依据；")
    doc.add_paragraph("（4）为系统测试与项目验收提供参考标准。")

    add_section_title(doc, "1.2  项目背景")
    doc.add_paragraph(
        "随着大语言模型技术的发展，传统问答系统已无法满足复杂任务处理与多轮智能交互需求。当前多数智能问答系统存在缺少统一交互支持、响应速度较慢、系统扩展性不足等问题。"
    )
    doc.add_paragraph(
        "为解决上述问题，本项目设计并实现AI Agent系统，采用前后端分离架构，支持Web端与命令行交互方式，满足教学实验与智能问答场景需求。"
    )

    add_section_title(doc, "1.3  项目目标")
    doc.add_paragraph("（1）实现具备智能对话能力的AI Agent系统；")
    doc.add_paragraph("（2）支持用户登录、会话管理与聊天记录保存；")
    doc.add_paragraph("（3）支持普通对话与流式对话功能；")
    doc.add_paragraph("（4）支持多种大语言模型服务切换；")
    doc.add_paragraph("（5）实现系统统计与运行状态检测；")
    doc.add_paragraph("（6）提供Docker化部署能力。")

    add_section_title(doc, "1.4  参考资料")
    refs = [
        "[1] 《软件工程》课程实验指导书",
        "[2] 实验二示例：需求说明书（面向对象方法学）",
        "[3] AI Agent项目源码",
        "[4] Spring Boot官方文档",
        "[5] React官方文档",
        "[6] UML建模相关资料",
    ]
    for r in refs:
        doc.add_paragraph(r)

    add_section_title(doc, "1.5  术语和缩写词")
    add_term_table(
        doc,
        [
            ("AI Agent", "智能代理系统"),
            ("JWT", "用户身份认证令牌"),
            ("SSE", "SSE（服务端流式推送技术）"),
            ("Session", "用户会话上下文"),
            ("RAG", "基于知识检索的增强生成技术"),
            ("CLI", "命令行交互客户端"),
        ],
    )

    add_section_title(doc, "1.6  需求获取方式")
    doc.add_paragraph("第一，对当前AI Agent项目已有功能进行整理；")
    doc.add_paragraph("第二，对智能问答系统典型业务流程进行分析；")
    doc.add_paragraph("第三，调研用户使用场景与交互需求；")
    doc.add_paragraph("第四，结合软件工程课程实验规范；")
    doc.add_paragraph("第五，考虑系统可扩展性与工程化要求。")

    # 2 需求分析
    add_section_title(doc, "2  需求分析")

    add_section_title(doc, "2.1  用户角色分析")
    role_table = doc.add_table(rows=4, cols=2)
    role_table.style = "Table Grid"
    role_table.rows[0].cells[0].text = "角色"
    role_table.rows[0].cells[1].text = "职责"
    role_table.rows[1].cells[0].text = "普通用户"
    role_table.rows[1].cells[1].text = "智能对话、会话管理、历史查看"
    role_table.rows[2].cells[0].text = "管理员"
    role_table.rows[2].cells[1].text = "管理系统状态、查看统计信息"
    role_table.rows[3].cells[0].text = "开发维护人员"
    role_table.rows[3].cells[1].text = "部署系统、维护模型服务"

    add_section_title(doc, "2.2  功能需求")

    add_section_title(doc, "2.2.1  系统总体业务流程")
    for s in [
        "（1）用户进入系统并完成登录认证；",
        "（2）系统验证身份并建立会话；",
        "（3）用户创建新聊天会话；",
        "（4）用户输入问题并发送请求；",
        "（5）系统调用模型服务生成回答；",
        "（6）系统返回结果并保存记录；",
        "（7）用户查看统计与导出数据。",
    ]:
        doc.add_paragraph(s)

    add_section_title(doc, "2.2.2  系统功能结构")
    for s in [
        "① 用户认证与授权模块",
        "② 会话管理模块",
        "③ 智能对话模块",
        "④ 模型服务管理模块",
        "⑤ 系统统计与运行监控模块",
        "⑥ Web前端交互模块",
        "⑦ CLI命令行交互模块",
    ]:
        doc.add_paragraph(s)

    add_section_title(doc, "2.3  用例模型")

    add_section_title(doc, "2.3.1  用户认证模块")
    add_usecase(
        doc,
        "2.3.1.1  用户注册",
        "普通用户",
        "无",
        "创建用户账号",
        ["用户输入邮箱与密码 → 系统校验 → 保存信息 → 返回成功。"],
        ["邮箱已存在、数据不完整则提示重新填写。"],
    )
    add_usecase(
        doc,
        "2.3.1.2  用户登录",
        "普通用户",
        "已注册",
        "生成登录状态",
        ["输入账号密码 → 验证身份 → 生成令牌 → 进入首页。"],
        ["密码错误、登录限流则提示失败。"],
    )

    add_section_title(doc, "2.3.2  会话管理模块")
    add_usecase(
        doc,
        "2.3.2.1  创建会话",
        "普通用户",
        "用户已登录",
        "创建聊天会话",
        ["用户输入会话名称、选择模型服务，系统创建并保存会话信息。"],
    )
    add_usecase(
        doc,
        "2.3.2.2  查看历史会话",
        "普通用户",
        "存在历史会话",
        "返回历史会话内容",
        ["系统校验用户权限后，返回对应会话历史记录。"],
    )

    add_section_title(doc, "2.3.3  智能对话模块")
    add_usecase(
        doc,
        "2.3.3.1  普通对话",
        "普通用户",
        "存在有效会话",
        "保存聊天记录",
        ["用户发送消息 → 系统调用模型服务 → 返回完整结果 → 保存聊天记录。"],
    )
    add_usecase(
        doc,
        "2.3.3.2  流式对话",
        "普通用户",
        "存在有效会话",
        "完成流式消息输出",
        ["建立SSE（服务端流式推送技术）连接 → 实时输出内容 → 完成后保存消息。"],
    )

    add_section_title(doc, "2.3.4  系统统计模块")
    add_usecase(
        doc,
        "2.3.4.1  查看系统运行状态",
        "管理员、开发维护人员",
        "系统已启动",
        "返回系统状态信息",
        ["检测数据库、缓存与模型服务状态并返回检测结果。"],
    )
    add_usecase(
        doc,
        "2.3.4.2  查看运行统计报告",
        "管理员、开发维护人员",
        "存在统计数据",
        "生成运行统计报告",
        ["系统按时间范围统计运行数据，并展示统计结果与运行报告。"],
    )

    add_section_title(doc, "2.4  数据需求")
    add_section_title(doc, "2.4.1  核心数据实体")
    doc.add_paragraph("User、Session、Message、ToolStat、ReleaseReport。")

    add_section_title(doc, "2.4.2  数据关系")
    for s in [
        "一个用户对应多个会话；",
        "一个会话包含多条消息；",
        "统计数据由用户操作行为生成；",
        "运行报告由系统自动生成。",
    ]:
        doc.add_paragraph(s)

    add_section_title(doc, "2.5  性能需求")
    for s in [
        "（1）普通聊天平均响应时间不超过3秒；",
        "（2）流式聊天支持连续输出；",
        "（3）系统状态检测响应时间不超过1秒；",
        "（4）支持多用户并发访问；",
        "（5）保证用户数据安全隔离。",
    ]:
        doc.add_paragraph(s)

    add_section_title(doc, "2.6  非功能需求")
    add_section_title(doc, "2.6.1  安全性")
    doc.add_paragraph("采用JWT认证机制，接口权限校验，密码加密存储，并具备登录限流功能。")
    add_section_title(doc, "2.6.2  可维护性")
    doc.add_paragraph("系统采用模块化与分层架构设计，模型服务支持扩展与替换，代码结构清晰。")
    add_section_title(doc, "2.6.3  可扩展性")
    doc.add_paragraph("支持新增模型服务、新增工具模块以及后续功能扩展。")
    add_section_title(doc, "2.6.4  易用性")
    doc.add_paragraph("系统界面简洁清晰，支持Web与CLI双端交互，操作便捷。")
    add_section_title(doc, "2.6.5  可部署性")
    doc.add_paragraph("支持Docker Compose一键部署，并支持脚本部署与环境变量配置。")

    # 3 运行环境
    add_section_title(doc, "3  运行环境")
    add_section_title(doc, "3.1  运行环境")
    for s in [
        "Web服务：Nginx",
        "后端服务：Spring Boot",
        "数据库：PostgreSQL",
        "缓存：Redis",
        "模型服务：大语言模型服务（OpenAI兼容模型、本地模型）",
    ]:
        doc.add_paragraph(s)

    add_section_title(doc, "3.2  开发环境")
    for s in [
        "操作系统：Windows 10",
        "开发工具：IntelliJ IDEA、VS Code",
        "构建工具：Maven、npm",
        "容器工具：Docker",
    ]:
        doc.add_paragraph(s)

    # 4 UML
    add_section_title(doc, "4  UML图清单")
    for s in [
        "图1-1 AI Agent系统总体活动图（描述用户从登录系统、创建会话到完成智能对话的整体业务流程）。",
        "图1-2 系统功能结构图（展示系统主要功能模块及模块之间关系）。",
        "图1-3 用户认证用例图（描述用户注册、登录与身份认证过程）。",
        "图1-4 会话管理用例图（描述会话创建、历史查看与会话管理流程）。",
        "图1-5 智能对话用例图（描述普通对话与流式对话功能流程）。",
        "图1-6 系统统计用例图（描述系统状态检测与统计报告功能）。",
        "图1-7 Web与CLI协同用例图（描述Web端与命令行客户端之间的协同交互方式）。",
    ]:
        doc.add_paragraph(s)

    add_section_title(doc, "UML 图示")
    add_uml_figure(doc, "图1_总体活动图.png", "图1-1 AI Agent系统总体活动图")
    add_uml_figure(doc, "图2_系统功能结构图.png", "图1-2 系统功能结构图")
    add_uml_figure(doc, "图3_认证与授权用例图.png", "图1-3 用户认证用例图")
    add_uml_figure(doc, "图4_会话管理用例图.png", "图1-4 会话管理用例图")
    add_uml_figure(doc, "图5_同步与流式对话用例图.png", "图1-5 智能对话用例图")
    add_uml_figure(doc, "图6_系统诊断与报表用例图.png", "图1-6 系统统计用例图")
    add_uml_figure(doc, "图7_Web与CLI多端协同用例图.png", "图1-7 Web与CLI协同用例图")

    # 5 结论
    add_section_title(doc, "5  实验结论")
    doc.add_paragraph(
        "本实验基于AI Agent系统完成需求分析与需求规格说明书编写，掌握了面向对象需求分析、用例建模、UML建模与工程化文档编写方法。"
    )
    doc.add_paragraph(
        "系统采用Spring Boot、React、PostgreSQL、Redis、Docker等技术实现，满足智能对话、会话管理、流式响应、模型切换等功能，符合课程实验要求，为后续设计、实现与测试奠定基础。"
    )

    # 强制正文小四宋体（简单兜底：遍历 runs）
    for para in doc.paragraphs:
        for run in para.runs:
            run.font.name = "宋体"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
            if run.font.size is None:
                run.font.size = Pt(12)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))


if __name__ == "__main__":
    out = ROOT / "实验二_AI Agent系统需求规格说明_软件工程实验报告_含UML图.docx"
    build_doc(out)
    print(out)
